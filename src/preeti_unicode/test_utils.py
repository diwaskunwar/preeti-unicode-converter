"""
Test utilities for the Preeti Unicode converter.

This module provides convenient test functions for users to quickly
test the package functionality with sample data.
"""

import os
from pathlib import Path
from typing import Optional, Union
import tempfile

from preeti_unicode.converter import convert_text
from preeti_unicode.file_converter import file_converter


def test(test_type: str = "string", input_data: Optional[str] = None, verbose: bool = True) -> Union[str, bool]:
    """
    Test the Preeti Unicode converter with different input types.
    
    This function provides a convenient way to test the package functionality
    with sample data or user-provided input.
    
    Args:
        test_type: Type of test to run ('string', 'txt', 'pdf', 'docx', 'all')
        input_data: Optional custom input data (for string test)
        verbose: Whether to print detailed output
        
    Returns:
        For 'string' test: converted text
        For file tests: True if successful, False otherwise
        For 'all' test: dictionary with results
        
    Examples:
        >>> from preeti_unicode import test
        >>> 
        >>> # Test string conversion
        >>> result = test("string", "g]kfn")
        >>> print(result)  # Output: नेपाल
        >>> 
        >>> # Test with sample files
        >>> test("txt")
        >>> test("pdf")
        >>> test("docx")
        >>> 
        >>> # Test all formats
        >>> test("all")
    """
    
    if test_type == "string":
        return _test_string_conversion(input_data, verbose)
    elif test_type == "txt":
        return _test_txt_conversion(verbose)
    elif test_type == "pdf":
        return _test_pdf_conversion(verbose)
    elif test_type == "docx":
        return _test_docx_conversion(verbose)
    elif test_type == "all":
        return _test_all_formats(verbose)
    else:
        if verbose:
            print(f"Unknown test type: {test_type}")
            print("Available types: 'string', 'txt', 'pdf', 'docx', 'all'")
        return False


def _test_string_conversion(input_data: Optional[str] = None, verbose: bool = True) -> str:
    """Test string conversion functionality."""
    
    # Use sample text if no input provided
    if input_data is None:
        input_data = "g]kfn Pp6f ;'Gb/ b]z xf] ."
    
    if verbose:
        print("=" * 50)
        print("STRING CONVERSION TEST")
        print("=" * 50)
        print(f"Input (Preeti): {input_data}")
    
    try:
        result = convert_text(input_data)
        
        if verbose:
            print(f"Output (Unicode): {result}")
            print("✓ String conversion test PASSED")
        
        return result
        
    except Exception as e:
        if verbose:
            print(f"✗ String conversion test FAILED: {e}")
        return ""


def _test_txt_conversion(verbose: bool = True) -> bool:
    """Test TXT file conversion functionality."""
    
    if verbose:
        print("=" * 50)
        print("TXT FILE CONVERSION TEST")
        print("=" * 50)
    
    try:
        # Create sample TXT file
        sample_text = """g]kfn Pp6f ;'Gb/ b]z xf] .
o; b]zsf] /fhwfgL sf7df8f}+ xf] .
g]kfnL efiff b]jgfu/L lnlkdf n]lvG5 .
o; b]zdf w]/} hfthfltsf dflg;x? a:5g\\ ."""
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write(sample_text)
            input_file = f.name
        
        # Create output file path
        output_file = input_file.replace('.txt', '_converted.txt')
        
        if verbose:
            print(f"Converting: {input_file}")
            print(f"Output: {output_file}")
        
        # Convert the file
        success = file_converter(input_file, 'txt', output_file, 'txt')
        
        if success and os.path.exists(output_file):
            # Read and display result
            with open(output_file, 'r', encoding='utf-8') as f:
                converted_content = f.read()
            
            if verbose:
                print("Converted content:")
                print(converted_content[:200] + "..." if len(converted_content) > 200 else converted_content)
                print("✓ TXT conversion test PASSED")
            
            # Clean up
            os.unlink(input_file)
            os.unlink(output_file)
            
            return True
        else:
            if verbose:
                print("✗ TXT conversion test FAILED: Conversion unsuccessful")
            
            # Clean up
            if os.path.exists(input_file):
                os.unlink(input_file)
            if os.path.exists(output_file):
                os.unlink(output_file)
            
            return False
            
    except Exception as e:
        if verbose:
            print(f"✗ TXT conversion test FAILED: {e}")
        return False


def _test_pdf_conversion(verbose: bool = True) -> bool:
    """Test PDF file conversion functionality."""
    
    if verbose:
        print("=" * 50)
        print("PDF FILE CONVERSION TEST")
        print("=" * 50)
    
    try:
        # Try to import required libraries
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4
        except ImportError:
            if verbose:
                print("✗ PDF test SKIPPED: reportlab not available")
            return False
        
        # Create sample PDF file
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            pdf_file = f.name
        
        # Create PDF with Preeti text
        c = canvas.Canvas(pdf_file, pagesize=A4)
        width, height = A4
        
        c.setFont('Helvetica', 12)
        y_position = height - 100
        
        preeti_lines = [
            "g]kfn Pp6f ;'Gb/ b]z xf] .",
            "o; b]zsf] /fhwfgL sf7df8f}+ xf] .",
            "g]kfnL efiff b]jgfu/L lnlkdf n]lvG5 ."
        ]
        
        for line in preeti_lines:
            c.drawString(100, y_position, line)
            y_position -= 30
        
        c.save()
        
        # Create output file path
        output_file = pdf_file.replace('.pdf', '_converted.txt')
        
        if verbose:
            print(f"Converting: {pdf_file}")
            print(f"Output: {output_file}")
        
        # Convert the file
        success = file_converter(pdf_file, 'pdf', output_file, 'txt')
        
        if success and os.path.exists(output_file):
            # Read and display result
            with open(output_file, 'r', encoding='utf-8') as f:
                converted_content = f.read()
            
            if verbose:
                print("Converted content:")
                print(converted_content[:200] + "..." if len(converted_content) > 200 else converted_content)
                print("✓ PDF conversion test PASSED")
            
            # Clean up
            os.unlink(pdf_file)
            os.unlink(output_file)
            
            return True
        else:
            if verbose:
                print("✗ PDF conversion test FAILED: Conversion unsuccessful")
            
            # Clean up
            if os.path.exists(pdf_file):
                os.unlink(pdf_file)
            if os.path.exists(output_file):
                os.unlink(output_file)
            
            return False
            
    except Exception as e:
        if verbose:
            print(f"✗ PDF conversion test FAILED: {e}")
        return False


def _test_docx_conversion(verbose: bool = True) -> bool:
    """Test DOCX file conversion functionality."""
    
    if verbose:
        print("=" * 50)
        print("DOCX FILE CONVERSION TEST")
        print("=" * 50)
    
    try:
        # Try to import required libraries
        try:
            from docx import Document
        except ImportError:
            if verbose:
                print("✗ DOCX test SKIPPED: python-docx not available")
            return False
        
        # Create sample DOCX file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            docx_file = f.name
        
        # Create DOCX with Preeti text
        doc = Document()
        doc.add_heading('g]kfnL kf7', 0)
        
        paragraphs = [
            "g]kfn Pp6f ;'Gb/ b]z xf] .",
            "o; b]zsf] /fhwfgL sf7df8f}+ xf] .",
            "g]kfnL efiff b]jgfu/L lnlkdf n]lvG5 ."
        ]
        
        for para_text in paragraphs:
            doc.add_paragraph(para_text)
        
        doc.save(docx_file)
        
        # Create output file path
        output_file = docx_file.replace('.docx', '_converted.txt')
        
        if verbose:
            print(f"Converting: {docx_file}")
            print(f"Output: {output_file}")
        
        # Convert the file
        success = file_converter(docx_file, 'docx', output_file, 'txt')
        
        if success and os.path.exists(output_file):
            # Read and display result
            with open(output_file, 'r', encoding='utf-8') as f:
                converted_content = f.read()
            
            if verbose:
                print("Converted content:")
                print(converted_content[:200] + "..." if len(converted_content) > 200 else converted_content)
                print("✓ DOCX conversion test PASSED")
            
            # Clean up
            os.unlink(docx_file)
            os.unlink(output_file)
            
            return True
        else:
            if verbose:
                print("✗ DOCX conversion test FAILED: Conversion unsuccessful")
            
            # Clean up
            if os.path.exists(docx_file):
                os.unlink(docx_file)
            if os.path.exists(output_file):
                os.unlink(output_file)
            
            return False
            
    except Exception as e:
        if verbose:
            print(f"✗ DOCX conversion test FAILED: {e}")
        return False


def _test_all_formats(verbose: bool = True) -> dict:
    """Test all supported formats."""
    
    if verbose:
        print("=" * 60)
        print("COMPREHENSIVE TEST - ALL FORMATS")
        print("=" * 60)
    
    results = {}
    
    # Test string conversion
    try:
        result = _test_string_conversion(verbose=verbose)
        results['string'] = bool(result)
    except Exception as e:
        results['string'] = False
        if verbose:
            print(f"String test error: {e}")
    
    if verbose:
        print()
    
    # Test TXT conversion
    results['txt'] = _test_txt_conversion(verbose=verbose)
    
    if verbose:
        print()
    
    # Test PDF conversion
    results['pdf'] = _test_pdf_conversion(verbose=verbose)
    
    if verbose:
        print()
    
    # Test DOCX conversion
    results['docx'] = _test_docx_conversion(verbose=verbose)
    
    if verbose:
        print()
        print("=" * 60)
        print("TEST SUMMARY")
        print("=" * 60)
        for test_type, success in results.items():
            status = "✓ PASSED" if success else "✗ FAILED"
            print(f"{test_type.upper()}: {status}")
        
        passed = sum(results.values())
        total = len(results)
        print(f"\nOverall: {passed}/{total} tests passed")
    
    return results


# Convenience functions for quick testing
def test_string(text: str = None) -> str:
    """Quick string conversion test."""
    return test("string", text, verbose=True)


def test_files() -> dict:
    """Quick test of all file formats."""
    return test("all", verbose=True)
