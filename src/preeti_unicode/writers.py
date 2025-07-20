"""
File writers for different output formats.

This module contains writer classes for creating output files in various formats
including PDF, DOCX, TXT, and HTML files with proper Unicode support.
"""

import os
from pathlib import Path
from typing import Dict, Any, List, Optional
from abc import ABC, abstractmethod

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.enums import TA_LEFT
except ImportError:
    SimpleDocTemplate = None

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    Document = None


class BaseWriter(ABC):
    """Abstract base class for file writers."""
    
    @abstractmethod
    def write(self, content: Dict[str, Any], file_path: Path, **kwargs) -> bool:
        """
        Write content to a file.
        
        Args:
            content: Content dictionary to write
            file_path: Path where to write the file
            **kwargs: Additional options
            
        Returns:
            True if successful, False otherwise
        """
        pass


class TXTWriter(BaseWriter):
    """Writer for plain text files."""
    
    def write(self, content: Dict[str, Any], file_path: Path, encoding: str = 'utf-8', **kwargs) -> bool:
        """
        Write content to a text file.
        
        Args:
            content: Content dictionary with 'text' key
            file_path: Path to the output text file
            encoding: Text encoding (default: utf-8)
            **kwargs: Additional options
            
        Returns:
            True if successful, False otherwise
        """
        try:
            text_content = content.get('text', '')
            
            # If content has pages, combine them
            if 'pages' in content:
                page_texts = []
                for i, page in enumerate(content['pages']):
                    if isinstance(page, dict) and 'text' in page:
                        page_texts.append(f"PAGE {i + 1}\n{'=' * 60}\n\n{page['text']}")
                    elif isinstance(page, str):
                        page_texts.append(f"PAGE {i + 1}\n{'=' * 60}\n\n{page}")
                text_content = '\n\n\n'.join(page_texts)
            
            with open(file_path, 'w', encoding=encoding) as file:
                file.write(text_content)
            
            return True
            
        except Exception as e:
            print(f"Error writing text file {file_path}: {e}")
            return False


class HTMLWriter(BaseWriter):
    """Writer for HTML files."""
    
    def write(self, content: Dict[str, Any], file_path: Path, **kwargs) -> bool:
        """
        Write content to an HTML file.
        
        Args:
            content: Content dictionary
            file_path: Path to the output HTML file
            **kwargs: Additional options
            
        Returns:
            True if successful, False otherwise
        """
        try:
            html_content = self._generate_html(content)
            
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(html_content)
            
            return True
            
        except Exception as e:
            print(f"Error writing HTML file {file_path}: {e}")
            return False
    
    def _generate_html(self, content: Dict[str, Any]) -> str:
        """Generate HTML content from the content dictionary."""
        html = """<!DOCTYPE html>
<html lang="ne">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Converted Document</title>
    <style>
        body {
            font-family: 'Noto Sans Devanagari', 'Arial Unicode MS', Arial, sans-serif;
            line-height: 1.6;
            margin: 40px;
            background-color: #f9f9f9;
        }
        .container {
            max-width: 800px;
            margin: 0 auto;
            background-color: white;
            padding: 40px;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }
        .page {
            margin-bottom: 40px;
            padding-bottom: 20px;
            border-bottom: 1px solid #eee;
        }
        .page:last-child {
            border-bottom: none;
        }
        .page-header {
            font-size: 18px;
            font-weight: bold;
            color: #333;
            margin-bottom: 20px;
            padding-bottom: 10px;
            border-bottom: 2px solid #007cba;
        }
        .content {
            font-size: 16px;
            line-height: 1.8;
            color: #444;
        }
        .paragraph {
            margin-bottom: 15px;
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>Converted Document</h1>
"""
        
        if 'pages' in content:
            for i, page in enumerate(content['pages']):
                html += f'        <div class="page">\n'
                html += f'            <div class="page-header">Page {i + 1}</div>\n'
                html += f'            <div class="content">\n'
                
                if isinstance(page, dict) and 'text' in page:
                    text = page['text']
                elif isinstance(page, str):
                    text = page
                else:
                    text = str(page)
                
                # Convert paragraphs to HTML
                paragraphs = text.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        html += f'                <div class="paragraph">{self._escape_html(para.strip())}</div>\n'
                
                html += f'            </div>\n'
                html += f'        </div>\n'
        else:
            # Single content
            text = content.get('text', '')
            html += f'        <div class="content">\n'
            paragraphs = text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    html += f'            <div class="paragraph">{self._escape_html(para.strip())}</div>\n'
            html += f'        </div>\n'
        
        html += """    </div>
</body>
</html>"""
        
        return html
    
    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters."""
        return (text.replace('&', '&amp;')
                   .replace('<', '&lt;')
                   .replace('>', '&gt;')
                   .replace('"', '&quot;')
                   .replace("'", '&#x27;')
                   .replace('\n', '<br>'))


class PDFWriter(BaseWriter):
    """Writer for PDF files."""
    
    def __init__(self):
        """Initialize PDF writer."""
        if SimpleDocTemplate is None:
            raise ImportError("reportlab is required for PDF writing. Install with: pip install reportlab")
    
    def write(self, content: Dict[str, Any], file_path: Path, **kwargs) -> bool:
        """
        Write content to a PDF file.
        
        Args:
            content: Content dictionary
            file_path: Path to the output PDF file
            **kwargs: Additional options
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Setup Unicode font
            font_name = self._setup_unicode_font()
            
            # Create PDF document
            doc = SimpleDocTemplate(str(file_path), pagesize=A4)
            story = []
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Create custom styles with Unicode font
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading1'],
                fontName=font_name,
                fontSize=16,
                spaceAfter=12,
                alignment=TA_LEFT,
                encoding='utf-8'
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=12,
                spaceAfter=6,
                alignment=TA_LEFT,
                encoding='utf-8'
            )
            
            if 'pages' in content:
                # Multi-page content
                for i, page in enumerate(content['pages']):
                    # Add page header
                    story.append(Paragraph(f"Page {i + 1}", heading_style))
                    story.append(Spacer(1, 0.2*inch))
                    
                    # Add page content
                    if isinstance(page, dict) and 'text' in page:
                        text = page['text']
                    elif isinstance(page, str):
                        text = page
                    else:
                        text = str(page)
                    
                    self._add_text_to_story(story, text, normal_style)
                    
                    # Add page break except for last page
                    if i < len(content['pages']) - 1:
                        story.append(PageBreak())
            else:
                # Single content
                text = content.get('text', '')
                self._add_text_to_story(story, text, normal_style)
            
            # Build PDF
            doc.build(story)
            return True
            
        except Exception as e:
            print(f"Error writing PDF file {file_path}: {e}")
            return False
    
    def _setup_unicode_font(self) -> str:
        """Setup Unicode font for proper Devanagari rendering."""
        try:
            # Try to use system fonts that support Devanagari
            font_paths = [
                '/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf',
                '/usr/share/fonts/truetype/noto/NotoSerifDevanagari-Regular.ttf',
                '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
                '/System/Library/Fonts/Arial.ttf',  # macOS
                'C:/Windows/Fonts/arial.ttf',  # Windows
            ]
            
            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont('DevanagariFont', font_path))
                        return 'DevanagariFont'
                    except Exception:
                        continue
            
            # Fallback to Helvetica
            return 'Helvetica'
            
        except Exception:
            return 'Helvetica'
    
    def _add_text_to_story(self, story, text: str, style):
        """Add text content to the PDF story."""
        if text.strip():
            paragraphs = text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    clean_para = para.replace('\n', ' ').strip()
                    clean_para = clean_para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    try:
                        story.append(Paragraph(clean_para, style))
                        story.append(Spacer(1, 0.1*inch))
                    except Exception:
                        # Fallback for problematic text
                        story.append(Paragraph(clean_para.encode('utf-8', 'replace').decode('utf-8'), style))
                        story.append(Spacer(1, 0.1*inch))


class DOCXWriter(BaseWriter):
    """Writer for DOCX files."""
    
    def __init__(self):
        """Initialize DOCX writer."""
        if Document is None:
            raise ImportError("python-docx is required for DOCX writing. Install with: pip install python-docx")
    
    def write(self, content: Dict[str, Any], file_path: Path, **kwargs) -> bool:
        """
        Write content to a DOCX file.
        
        Args:
            content: Content dictionary
            file_path: Path to the output DOCX file
            **kwargs: Additional options
            
        Returns:
            True if successful, False otherwise
        """
        try:
            doc = Document()
            
            if 'pages' in content:
                # Multi-page content
                for i, page in enumerate(content['pages']):
                    # Add page header
                    heading = doc.add_heading(f'Page {i + 1}', level=1)
                    
                    # Add page content
                    if isinstance(page, dict) and 'text' in page:
                        text = page['text']
                    elif isinstance(page, str):
                        text = page
                    else:
                        text = str(page)
                    
                    self._add_text_to_doc(doc, text)
                    
                    # Add page break except for last page
                    if i < len(content['pages']) - 1:
                        doc.add_page_break()
            else:
                # Single content
                text = content.get('text', '')
                self._add_text_to_doc(doc, text)
            
            doc.save(str(file_path))
            return True
            
        except Exception as e:
            print(f"Error writing DOCX file {file_path}: {e}")
            return False
    
    def _add_text_to_doc(self, doc, text: str):
        """Add text content to the DOCX document."""
        if text.strip():
            paragraphs = text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
